"""
Embedding Service — Processes uploaded documents and indexes them into FAISS.

Pipeline:
1. Read document (PDF/DOCX/TXT)
2. Split into chunks using RecursiveCharacterTextSplitter
3. Generate embeddings with BAAI/bge-small-en-v1.5
4. Upsert into FAISS vector store
5. Save updated index to disk
"""
import os
import logging
import threading
from pathlib import Path
from typing import List

from django.conf import settings

logger = logging.getLogger(__name__)

_index_write_lock = threading.Lock()


def extract_text_from_file(file_path: str, file_type: str) -> str:
    """Extract raw text content from PDF, DOCX, or TXT files."""
    file_path = str(file_path)

    if file_type == 'pdf':
        text_parts = []
        try:
            try:
                import pypdf
                reader = pypdf.PdfReader(file_path)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                return '\n\n'.join(text_parts)
            except ImportError:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                return '\n\n'.join(text_parts)
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise ValueError(f"Failed to read PDF: {e}")

    elif file_type == 'docx':
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return '\n\n'.join(paragraphs)
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise ValueError(f"Failed to read DOCX: {e}")

    elif file_type == 'txt':
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            logger.error(f"TXT extraction error: {e}")
            raise ValueError(f"Failed to read TXT: {e}")

    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def chunk_text(text: str, source_name: str) -> List:
    """Split text into overlapping chunks with metadata."""
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter
    except ImportError:
        from langchain.text_splitter import RecursiveCharacterTextSplitter

    try:
        from langchain_core.documents import Document as LCDocument
    except ImportError:
        from langchain.schema import Document as LCDocument

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=100,
        length_function=len,
        separators=['\n\n', '\n', '. ', '! ', '? ', ' ', ''],
    )

    chunks = splitter.split_text(text)
    documents = [
        LCDocument(
            page_content=chunk,
            metadata={'source': source_name}
        )
        for chunk in chunks if chunk.strip()
    ]
    return documents


def index_document(document_record) -> int:
    """
    Full pipeline: extract text → chunk → embed → store in FAISS.
    Uses thread lock to ensure thread safety during FAISS disk updates.
    """
    try:
        from langchain_huggingface import HuggingFaceEmbeddings
    except ImportError:
        from langchain_community.embeddings import HuggingFaceEmbeddings
    from langchain_community.vectorstores import FAISS

    store_path = str(settings.VECTOR_STORE_PATH)
    Path(store_path).parent.mkdir(parents=True, exist_ok=True)

    file_path = document_record.file.path
    file_type = document_record.file_type.lower()
    source_name = document_record.title

    logger.info(f"Starting indexing for: {source_name}")

    text = extract_text_from_file(file_path, file_type)
    if not text.strip():
        raise ValueError("Document appears to be empty or unreadable.")

    chunks = chunk_text(text, source_name)
    if not chunks:
        raise ValueError("No text chunks were extracted from the document.")

    logger.info(f"Created {len(chunks)} chunks from {source_name}")

    embeddings = HuggingFaceEmbeddings(
        model_name="BAAI/bge-small-en-v1.5",
        model_kwargs={'device': 'cpu'},
        encode_kwargs={'normalize_embeddings': True}
    )

    index_file = f"{store_path}.faiss"

    with _index_write_lock:
        if os.path.exists(index_file) or os.path.exists(store_path):
            try:
                existing_vs = FAISS.load_local(
                    store_path,
                    embeddings,
                    allow_dangerous_deserialization=True
                )
                existing_vs.add_documents(chunks)
                existing_vs.save_local(store_path)
                logger.info(f"Merged {len(chunks)} chunks into existing FAISS index.")
            except Exception as e:
                logger.warning(f"Could not merge with existing index ({e}), creating new.")
                vs = FAISS.from_documents(chunks, embeddings)
                vs.save_local(store_path)
        else:
            vs = FAISS.from_documents(chunks, embeddings)
            vs.save_local(store_path)
            logger.info(f"Created new FAISS index with {len(chunks)} chunks.")

    from chatbot.ai_service import reload_vector_store
    reload_vector_store()

    return len(chunks)


def delete_document_from_index(document_title: str):
    """
    Remove a document's chunks from FAISS by source metadata.
    Uses thread lock to ensure safe rebuild.
    """
    try:
        from langchain_huggingface import HuggingFaceEmbeddings
    except ImportError:
        from langchain_community.embeddings import HuggingFaceEmbeddings
    from langchain_community.vectorstores import FAISS

    store_path = str(settings.VECTOR_STORE_PATH)
    index_file = f"{store_path}.faiss"

    if not (os.path.exists(index_file) or os.path.exists(store_path)):
        return

    with _index_write_lock:
        try:
            embeddings = HuggingFaceEmbeddings(
                model_name="BAAI/bge-small-en-v1.5",
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': True}
            )
            vs = FAISS.load_local(store_path, embeddings, allow_dangerous_deserialization=True)

            all_docs = []
            for doc_id, doc in vs.docstore._dict.items():
                if doc.metadata.get('source') != document_title:
                    all_docs.append(doc)

            if all_docs:
                new_vs = FAISS.from_documents(all_docs, embeddings)
                new_vs.save_local(store_path)
            else:
                import shutil
                if os.path.isdir(store_path):
                    shutil.rmtree(store_path)
                else:
                    for ext in ['.faiss', '.pkl']:
                        path = store_path + ext
                        if os.path.exists(path):
                            os.remove(path)

            from chatbot.ai_service import reload_vector_store
            reload_vector_store()
            logger.info(f"Removed chunks for '{document_title}' from FAISS index.")
        except Exception as e:
            logger.error(f"Error deleting document from index: {e}")


def index_url_content(url: str, title: str = '', user = None):
    """
    Fetch webpage HTML, extract clean text content, chunk it, and index into FAISS.
    """
    import urllib.request
    import re
    from html import unescape
    from .models import Document

    req = urllib.request.Request(
        url,
        headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) SupportGenie-RAG/1.0'}
    )

    try:
        with urllib.request.urlopen(req, timeout=15) as resp:
            html = resp.read().decode('utf-8', errors='ignore')
    except Exception as e:
        logger.error(f"Failed to fetch URL {url}: {e}")
        raise ValueError(f"Could not access webpage: {str(e)}")

    # Extract title from HTML if not provided
    if not title:
        title_match = re.search(r'<title[^>]*>(.*?)</title>', html, re.IGNORECASE | re.DOTALL)
        if title_match:
            title = unescape(title_match.group(1)).strip()
        else:
            title = url

    # Remove script, style, and metadata tags
    cleaned_html = re.sub(r'<(script|style|header|footer|nav)[^>]*>.*?</\1>', '', html, flags=re.IGNORECASE | re.DOTALL)
    # Strip HTML tags
    text = re.sub(r'<[^>]+>', ' ', cleaned_html)
    # Unescape HTML entities and normalize whitespace
    text = unescape(text)
    text = re.sub(r'\s+', ' ', text).strip()

    if len(text) < 100:
        raise ValueError("Webpage content was too short or contained no extractable text.")

    source_name = f"Web: {title}"

    # Create Document record
    doc_record = Document.objects.create(
        title=title,
        file_type='url',
        file_size=len(text.encode('utf-8')),
        uploaded_by=user,
        status='processing',
        chunk_count=0
    )

    chunks = chunk_text(text, source_name)
    if not chunks:
        doc_record.status = 'failed'
        doc_record.error_message = 'No readable text chunks extracted.'
        doc_record.save()
        raise ValueError("Could not extract readable text chunks from this webpage.")

    store_path = str(settings.VECTOR_STORE_PATH)
    Path(store_path).parent.mkdir(parents=True, exist_ok=True)

    try:
        from langchain_huggingface import HuggingFaceEmbeddings
    except ImportError:
        from langchain_community.embeddings import HuggingFaceEmbeddings
    from langchain_community.vectorstores import FAISS

    embeddings = HuggingFaceEmbeddings(
        model_name="BAAI/bge-small-en-v1.5",
        model_kwargs={'device': 'cpu'},
        encode_kwargs={'normalize_embeddings': True}
    )

    index_file = f"{store_path}.faiss"

    with _index_write_lock:
        if os.path.exists(index_file) or os.path.exists(store_path):
            try:
                existing_vs = FAISS.load_local(
                    store_path,
                    embeddings,
                    allow_dangerous_deserialization=True
                )
                existing_vs.add_documents(chunks)
                existing_vs.save_local(store_path)
            except Exception as e:
                vs = FAISS.from_documents(chunks, embeddings)
                vs.save_local(store_path)
        else:
            vs = FAISS.from_documents(chunks, embeddings)
            vs.save_local(store_path)

    doc_record.status = 'indexed'
    doc_record.chunk_count = len(chunks)
    doc_record.save()

    from chatbot.ai_service import reload_vector_store
    reload_vector_store()

    return doc_record
